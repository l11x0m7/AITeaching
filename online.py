import flask

from flask import Flask
# app = Flask(__name__)

# @app.route('/')
# def hello_world():
#    return 'Hello World'

# if __name__ == '__main__':
#    app.run(host='0.0.0.0', port=8080)



#!/usr/bin/env python
# coding=utf-8 
import bottle
from bottle import get, post, request, route, run, template, static_file
import threading
import json
import numpy as np

from time import sleep

import os
import sys
import torch
from fairseq.data.data_utils import collate_tokens
from fairseq.models.roberta import RobertaModel
import re
import uuid

from qingstor.sdk.service.qingstor import QingStor
from qingstor.sdk.config import Config

from docx import Document
from docx.shared import RGBColor
from curses import ascii

ACCESS_KEY_ID = 'KPXLUFSRVNVNZGFCEPDT'
SECRET_ACCESS_KEY = '9RW7JW2RsIDmArXSdeHhCjYt7A9vHPs6LBT8zSEp'
prefix = 'https://mrc-lxm.pek3b.qingstor.com/'
tmp_path = '/tmp/mrc'

config = Config(ACCESS_KEY_ID, SECRET_ACCESS_KEY)
qingstor = QingStor(config)
bucket = qingstor.Bucket('mrc-lxm', 'pek3b')


# -- translate --
from translate import translate

def clear_unascii(s):
    return ''.join([c for c in s if ascii.isascii(c)])


'''
This file is taken and modified from R-Net by Minsangkim142
https://github.com/minsangkim142/R-net
'''

query = []
response = []
app = bottle.Bottle()


@app.route('/img/<filename:re:.*\.png|.*\.jpg>')
def server_static(filename):
    return static_file(filename, root='./trendsetter/img/')
@app.route('/img/portfolio/<filename:re:.*\.png|.*\.jpg>')
def server_static(filename):
    return static_file(filename, root='./trendsetter/img/portfolio/')
@app.route('/css/<filename:re:.*\.css>')
def server_static(filename):
    return static_file(filename, root='./trendsetter/css/')
@app.route('/js/<filename:re:.*\.js>')
def server_static(filename):
    return static_file(filename, root='./trendsetter/js/')
@app.route('/js/vendor/<filename:re:.*\.js>')
def server_static(filename):
    return static_file(filename, root='./trendsetter/js/vendor/')
@app.route('/<filename:re:.*\.html>')
def server_static(filename):
    return static_file(filename, root='./trendsetter/')


@app.get("/")
def home():
    # with open('demo.html', 'r') as fl:
    with open('trendsetter/index.html', 'r') as fl:
        html = fl.read()
        return html

@app.post('/answer')
def answer():
    passage = bottle.request.json['passage']
    question = bottle.request.json['question']
    print("received question: {}".format(question))
    print("received passage: {}".format(passage))
    # if not passage or not question:
    #     exit()
    global query, response
    query = [question, passage]
    while not response:
        sleep(0.1)
    print("received response: {}".format(str(response)))
    response_ = {"answer": response[0], "url": response[1], "detail": response[2]}
    response = []
    return response_


class Demo(object):
    def __init__(self, model):
        run_event = threading.Event()
        run_event.set()
        threading.Thread(target=self.demo_backend, args = [model, run_event]).start()
        app.run(port=80, host='0.0.0.0', debug=False)
        try:
            while 1:
                sleep(.1)
        except KeyboardInterrupt:
            print("Closing server...")
            run_event.clear()

    def demo_backend(self, model, run_event):
        global query, response

        while run_event.is_set():
            sleep(0.1)
            if query:
                fname = 'mrc_{}.docx'.format(uuid.uuid1())
                if not os.path.exists(tmp_path):
                    os.mkdir(tmp_path)
                # response = '答案:' + '计算中，请稍等...'
                qas, context = query[0], query[1]
                qas, context = preprocess_query_and_doc(qas, context)
                all_qas, ori_qas, options = get_format_output(qas, context)
                if debug:
                    response = []
                    # with open(os.path.join(tmp_path, fname), 'w') as fw:
                    #     fw.write('\n\n'.join([context, '\n\n'.join(['\n'.join(_) for _ in all_qas])]))
                    # extract_word_file([context for _ in range(len(ori_qas))], 
                    #                    [[1 for _ in range(len(context.split(' ')))] for __ in range(len(ori_qas))], 
                    #                    ori_qas, ['A' for _ in range(len(ori_qas))], [0.9 for _ in range(len(ori_qas))], fname)
                    details = {"doc": ["This is the passage.", "Sent2.", "Sent3."], "trans": ["这是一篇文章。", "句子2。", "句子3。"], "qas":[("q1", ["op11", "op12", "op13", "op14"], [0, 1, 3]), ("q2", ["op21", "op22", "op23", "op24"], [0, 1, 2])]}
                    
                    
                    response = ['答案:' + 'A', '{}{}'.format(prefix, fname), details]
                    query = []
                else:
                    response = get_MRC_answer(fname, context, all_qas, ori_qas, options)
                    query = []


def preprocess_query_and_doc(query, context):
    if query.startswith('问题和选项如:'):
        query = query[8:]
    if context.startswith('文章如:'):
        context = context[5:]
    context = context.replace("''", '" ').replace("``", '" ')
    qas = query
    qas = qas.replace('。', '.').replace('，', ',').replace('？', 
				'?').replace('！', '!').replace('（', '(').replace('）', 
				')').replace('‘', '\'').replace('’', '\'').replace('“', '"').replace('”', '"')
    context = context.replace('。', '.').replace('，', ',').replace('？', 
				'?').replace('！', '!').replace('（', '(').replace('）', 
				')').replace('‘', '\'').replace('’', '\'').replace('“', '"').replace('”', '"')
    qas = clear_unascii(qas)
    context = clear_unascii(context)
    return qas, context


def get_format_output(qas, context):
    flag = True
    q = ""
    all_qas = []
    ori_qas = []
    options = []
    for _ in qas.strip().split('\n'):
        if flag:
            q = _.strip()
            flag = False
            continue
        if _ == "":
            if len(options) == 1:
                options = options[0].replace('\t', ' ').strip().split(' ')
                options = [_ for _ in options if len(_) > 0]
            ori_qas.append('\n'.join([q] + options))
            options = [re.sub('^[ABCD]\.', '', o).strip() for o in options]
            options = [re.sub('^[ABCD] ', '', o).strip() for o in options]
            all_qas.append([q + ' ' + o for o in options])
            flag = True
            options = []
            continue
        options.append(_)
    if options:
        if len(options) == 1:
            options = re.split('[A-Z][\. ]', options[0].replace('\t', ' ').strip())
            options = [_.strip() for _ in options if len(_) > 0]
        ori_qas.append('\n'.join([q] + options))
        options = [re.sub('^[A-Z]\.', '', o).strip() for o in options]
        options = [re.sub('^[A-Z] ', '', o).strip() for o in options]
        all_qas.append([q + ' ' + o for o in options])
    print(all_qas)
    return all_qas, ori_qas, options


def get_MRC_answer(fname, context, all_qas, ori_qas, options):
    ans = ""
    answers = []
    confidents = []
    docs = []
    weights = []
    for qas in all_qas:
        ts = []
        for qa in qas:
            inp = model.encode(qa, context)
            ts.append(inp)
        batch = collate_tokens(ts, pad_idx=1)
        # print(model.extract_features_aligned_to_words(qa))
        logits, last_attn = model.predict('sentence_classification_head', batch, return_logits=True)
        logits = torch.nn.functional.softmax(logits.squeeze())
        print(last_attn.shape)
        logits = logits.tolist()
        logits = np.asarray(logits).flatten()
        print(logits)
        answer = np.argmax(logits)
        confident = logits[answer]
        print(torch.max(last_attn[answer, 0, :]))
        print(torch.sum(last_attn[answer, 0, :]))
        toks, attns = model.extract_attention_to_words(qas[answer], context, last_attn[answer, 0, :].squeeze())
        attns = attns.tolist()
        ans += chr(ord('A') + answer)
        answers.append(chr(ord('A') + answer))
        confidents.append(confident)
        docs.append(toks)
        weights.append(attns)
    extract_word_file(docs, weights, ori_qas, answers, confidents, fname)
    # response = []
    response = ['答案:' + ans, '{}{}'.format(prefix, fname)]
    return response

def cut_sent(para):
    para = re.sub('([\.!\?] ?)([^\'"]) ?', r"\1@@@@@\2", para)  # 单字符断句符
    para = re.sub('(\.{3,} ?)([^\'"]) ?', r"\1@@@@@\2", para)  # 英文省略号
    para = re.sub('(\…{2,} ?)([^\'"]) ?', r"\1@@@@@\2", para)  # 中文省略号
    para = re.sub('([.!\?]["\'] ?)([^,.!\?] ?)', r'\1@@@@@\2', para)
    para = para.rstrip(r'@@@@@')  # 段尾如果有多余的\n就去掉它
    return para.split(r'@@@@@')


def extract_word_file(docs, weights, ori_qas, answers, confidents, fname):
    # with open(os.path.join(tmp_path, fname), 'w') as fw:
    document = Document()
    p = document.add_paragraph('注: 信息关联度(从强到弱)为`红->橙->棕->绿->黑`\n\n\n')
    count = 1
    for doc, weight, ori_qa, answer, conf in zip(docs, weights, ori_qas, answers, confidents):
        weight = [_ for i, _ in enumerate(weight) if doc[i] != '' and doc[i] != ' ']
        doc = [_ for _ in doc if _ != '' and _ != ' ']
        run = p.add_run('题目{}:\n\n'.format(count))
        # weight = ' '.join(list(map(str, weight)))
        conf = str(conf)
        # print(doc)
        # print(weight)
        # print(ori_qa)
        # print(answer)
        # print(conf)
        #add_run在同一段添加内容
        run = p.add_run('文章:\n')
        doc = '#####'.join(doc)
        doc_sent = cut_sent(doc)
        trans_doc_sent = translate('\n'.join([''.join(_.split('#####')) for _ in doc_sent]))
        w_i = 0
        total_sent = len(doc_sent)
        colors = []
        for sent in doc_sent:
            sent = sent.split('#####')
            total_score = 0.
            for w in sent:
                if not w:
                    continue
                if w_i >= len(weight):
                    break
                wei= weight[w_i]
                w_i += 1
                if wei > 10. / len(weight):
                    total_score += 8
                elif wei > 6. / len(weight):
                    total_score += 6
                elif wei > 3. / len(weight):
                    total_score += 4
                elif wei > 1.5 / len(weight):
                    total_score += 2
            if total_score >= 14:
                color = RGBColor(255,0,0)
            elif total_score >= 10:
                color = RGBColor(255,165,0)
            elif total_score >= 8:
                color = RGBColor(165,42,42)
            elif total_score >= 6:
                color = RGBColor(0,128,0)
            else:
                color = RGBColor(0,0,0)
            colors.append(color)
            for w in sent:
                run = p.add_run(w)
                run.font.color.rgb = color
        run = p.add_run('\n翻译:\n')
        for _, trans_sent in enumerate(trans_doc_sent):
            for w in trans_sent:
                run = p.add_run(w)
                run.font.color.rgb = colors[_]
 
            # for w, wei in zip(doc, weight):
#                 if w == '\n':
#                     continue
#                 if wei > 10. / len(weight):
#                     run.font.color.rgb = RGBColor(255,0,0)
#                 elif wei > 6. / len(weight):
#                     run.font.color.rgb = RGBColor(255,165,0)
#                 elif wei > 3. / len(weight):
#                     run.font.color.rgb = RGBColor(165,42,42)
#                 elif wei > 1.5 / len(weight):
#                     run.font.color.rgb = RGBColor(0,128,0)
#                 else:
#                     run.font.color.rgb = RGBColor(0,0,0)
        # doc = ' '.join(doc)
        # doc = doc.replace(' \n ', '\n')
        # run = p.add_run('文章:\n' + doc + '\n\n')
        run = p.add_run('\n\n问题:\n' + ori_qa + '\n\n')
        run = p.add_run("最终答案: " + answer + '\n\n')
        run = p.add_run('置信度: ' + conf)
        run = p.add_run('\n' * 5)
        count += 1
            # fw.write('\n\n'.join(['注:信息关联度(从强到弱):红->橙->黄->黑\n', '文章:\n' + doc, weight, '问题:\n' + ori_qa, "最终答案:" + answer, '置信度:' + conf]) + '\n' * 5)
    document.save(os.path.join(tmp_path, fname))

    with open(os.path.join(tmp_path, fname), 'rb') as f:
        output = bucket.put_object(
            fname, body=f
        )
        if output.status_code == 201:
            print('Upload Success')
        else:
            print('Upload Failed')
    os.remove(os.path.join(tmp_path, fname))
        


if __name__ == '__main__':
    model = ""
    debug = True
    if not debug:
        roberta = RobertaModel.from_pretrained('checkpoints/', checkpoint_file='ck.pt', data_name_or_path='data/processed_RACE/')
        roberta.eval()
        model = roberta
    demo = Demo(model)
